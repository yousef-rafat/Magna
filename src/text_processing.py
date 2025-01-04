import os
import fitz
import requests
from io import BytesIO
from docx import Document

def get_pdf_text(file_path_or_url): 

    """ Reads the text content of a pdf file into a text variable. Works with local and internet-hosted documents. """

    if file_path_or_url.startswith('http'):

        response = requests.get(file_path_or_url)
        response.raise_for_status()

        pdf_file = BytesIO(response.content)
        doc = fitz.open("pdf", pdf_file.read())

    else:
        doc = fitz.open(pdf_file)

    text = ""
    for page in doc:
        text += page.get_text() + "\n"

    return text

class TextSplitter():
    
    """ A replacement for Langchain class for more efficient software 
        Used for splitting text into chunks with overlap between them
    """
    
    def __init__(self, chunk_size = 1280, chunk_overlap = 320, add_start_index = False):
        self.chunk_size = chunk_size
        self.overlap = chunk_size - chunk_overlap
        self.add_index = add_start_index

    def split_text(self, text):
        chunks = []
        start = 0

        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]
            if self.add_index:
                chunks.append((chunk, start))
            else:
                chunks.append(chunk)

            start += self.overlap

        return chunks
    

def get_text_from_docx(file_path_or_url):
    
    """ Get all the text from a docx file. Works with local and internet-hosted documents. """

    if file_path_or_url.startswith('http'):
        
        response = requests.get(file_path_or_url)
        response.raise_for_status()
        
        doc_file = BytesIO(response.content)
        document = Document(doc_file)

    else:
        document = Document(file_path_or_url)

    texts = []
    for paragraph in document.paragraphs:
        texts.append(paragraph.text)

    return "\n".join(texts)

def read_text_file(file_path):
    """ Reads a text file from disk """

    with open(file_path, "r") as file:
        texts = [line.strip() for line in file]        

    return texts


def find_files_from_special_folders(file_types=("pdf", "docx", "txt")):

    """

    Finds and prints files of specified types in the 'Documents' and 'Downloads' folders of the user.
    
    Parameters:
        file_types (tuple): A tuple of file extensions to search for (e.g., ("pdf", "docx")).

    """
    home_directory = os.path.expanduser("~")
    special_folders = ["Documents", "Downloads"]

    for folder in special_folders:
        target_directory = os.path.join(home_directory, folder)

        if os.path.exists(target_directory):
            print(f"\nSearching for files in: {target_directory}")

            for root, _, files in os.walk(target_directory):
                for file in files:
                    if file.lower().endswith(tuple(f".{ext.lower()}" for ext in file_types)):
                        file_path = os.path.join(root, file)
                        print(file_path + '\n')
        else:
            print(f"\nDirectory does not exist: {target_directory}")
